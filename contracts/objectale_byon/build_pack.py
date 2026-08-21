#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_pack.py — the objectale -> byon contract pack, German and English, from one source.

    python contracts/objectale_byon/build_pack.py [--out DIR]

ONE COMMAND, TEN FILES. Five agreements (reseller, white-label/OEM, mutual NDA, SLA, AVV) in two
languages. The content lives in doc_*.py as parallel block lists; this file renders them and
refuses to write anything until the two languages agree structurally.

THE GATE IS THE POINT. A bilingual contract set fails in one specific way: somebody edits the
German, the English keeps the old clause numbering, and six months later the two versions say
different things about liability. So before rendering, every document is checked clause by clause:
same block types in the same order, same clause numbers, same table shapes. That is a structural
property, it is cheap to assert, and it is the only thing standing between this pack and the
failure mode that makes a bilingual contract worthless.

It also counts the bracketed fields that a human still has to fill, and prints them, because an
unfilled placeholder in a signed contract is worse than a blank page.
"""
import argparse
import os
import re
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
if HERE not in sys.path:
    sys.path.insert(0, HERE)

import common as C                                                       # noqa: E402
import doc_distribution                                                  # noqa: E402
import doc_reseller                                                      # noqa: E402
import doc_deed                                                          # noqa: E402
import doc_eula                                                          # noqa: E402
import doc_oem                                                           # noqa: E402
import doc_nda                                                           # noqa: E402
import doc_sla                                                           # noqa: E402
import doc_dpa                                                           # noqa: E402

# CHAIN ORDER, and the numbering is the reading order for a customer's counsel: who owns it, who
# may sell it, who did sell it, and what the buyer signs up to.
DOCS = [
    ("00_Distributionsvertrag", "00_Distribution_Agreement", doc_distribution),
    ("01_Wiederverkaeufer_Rahmenvertrag", "01_Reseller_Framework_Agreement", doc_reseller),
    ("02_White_Label_OEM_Vertrag", "02_White_Label_OEM_Agreement", doc_oem),
    ("03_Gegenseitige_Geheimhaltungsvereinbarung", "03_Mutual_NDA", doc_nda),
    ("04_Service_Level_Agreement", "04_Service_Level_Agreement", doc_sla),
    ("05_Auftragsverarbeitungsvertrag", "05_Data_Processing_Agreement", doc_dpa),
    ("06_Durchgriffs_und_Eintrittsvereinbarung", "06_Flow_Down_and_Step_In_Deed", doc_deed),
    ("07_Endkundenbedingungen", "07_End_User_Terms", doc_eula),
]


# --------------------------------------------------------------------------- the gate
_SCHED = re.compile(r"^\s*(?:Schedule|Anlage|Annex|Anhang|Appendix)\s+([0-9A-Z]+)\b", re.I)
_CLAUSE = re.compile(r"^\s*([0-9]+(?:\.[0-9]+)*)\.")


def _ident(head):
    """The IDENTIFIER of a heading, with the label word removed.

    "Schedule 1" and "Anlage 1" are the same schedule; comparing the whole string made the gate
    fail on a correct translation, which is the fastest way to teach somebody to pass --force.
    A heading with no identifier (Background, Signatures) compares as its own kind, so those still
    have to appear in the same position in both languages.
    """
    m = _SCHED.match(head)
    if m:
        return "sched:" + m.group(1).upper()
    m = _CLAUSE.match(head)
    if m:
        return "cl:" + m.group(1)
    return "unnumbered"


def shape(blocks):
    """The structure of a document, with none of its words.

    Two languages of the same contract must produce an identical shape. Clause NUMBERS are part of
    it (a cross-reference to clause 17 has to point at liability in both languages) and table
    DIMENSIONS are part of it (a price schedule that gained a row in German only is a commercial
    dispute waiting to happen). The prose itself is deliberately not compared.
    """
    out = []
    for b in blocks:
        kind = b[0]
        if kind == "h2":
            out.append(("h2", _ident(b[1])))
        elif kind == "num":
            out.append(("num", b[1]))
        elif kind == "table":
            out.append(("table", len(b[1]), len(b[2]),
                        tuple(len(r) for r in b[2])))
        else:
            out.append((kind,))
    return out


def check(en, de, name, fails):
    a, b = shape(en), shape(de)
    if a == b:
        return
    fails.append(name)
    print("  [X] %s: the two languages are not parallel" % name)
    for i in range(max(len(a), len(b))):
        x = a[i] if i < len(a) else None
        y = b[i] if i < len(b) else None
        if x != y:
            print("      block %d   EN %-40s   DE %s" % (i, x, y))
            if len([1 for j in range(i, min(i + 4, max(len(a), len(b))))]) >= 3:
                break


def placeholders(blocks):
    found = []
    for b in blocks:
        for part in b[1:]:
            for s in _strings(part):
                found += re.findall(r"\[[^\[\]]{2,80}\]", s)
    return found


def _strings(x):
    if isinstance(x, str):
        yield x
    elif isinstance(x, (list, tuple)):
        for i in x:
            for s in _strings(i):
                yield s


# --------------------------------------------------------------------------- rendering
def render(blocks, path, lang):
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    d = Document()
    s = d.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.3)
    s.top_margin = s.bottom_margin = Cm(2.1)
    n = d.styles["Normal"]
    n.font.name = C.FONT_BODY
    n.font.size = Pt(10)
    n.font.color.rgb = RGBColor.from_string(C.INK)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.12

    def para(text="", size=10, bold=False, colour=C.INK, face=C.FONT_BODY,
             before=0, after=6, indent=0.0, italic=False):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        if text:
            r = p.add_run(text)
            r.font.name = face
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.italic = italic
            r.font.color.rgb = RGBColor.from_string(colour)
        return p

    def shade(cell, hexv):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexv)
        cell._tc.get_or_add_tcPr().append(el)

    for b in blocks:
        k = b[0]
        if k == "h1":
            para(b[1], size=19, face=C.FONT_HEAD, colour=C.HEAD, after=2)
        elif k == "meta":
            para(b[1], size=8.5, colour=C.MUTED, after=14)
        elif k == "h2":
            para(b[1], size=12, face=C.FONT_HEAD, colour=C.HEAD, before=15, after=5)
        elif k == "h3":
            para(b[1], size=10.5, bold=True, colour=C.ACCENT, before=9, after=3)
        elif k == "p":
            para(b[1])
        elif k == "num":
            p = d.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(b[1] + "  ")
            r.font.name = C.FONT_BODY
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(C.ACCENT)
            r2 = p.add_run(b[2])
            r2.font.name = C.FONT_BODY
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor.from_string(C.INK)
        elif k == "bullet":
            para("•  " + b[1], indent=0.5, after=3)
        elif k == "note":
            p = para(b[1], size=9, colour=C.MUTED, italic=True, before=8, after=10, indent=0.3)
            p.paragraph_format.right_indent = Cm(0.3)
        elif k == "table":
            head, rows = b[1], b[2]
            t = d.add_table(rows=len(rows) + (1 if head else 0), cols=len(rows[0]))
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            t.autofit = True
            ri = 0
            if head:
                for ci, h in enumerate(head):
                    cell = t.rows[0].cells[ci]
                    cell.text = ""
                    pr = cell.paragraphs[0].add_run(h)
                    pr.font.name = C.FONT_BODY
                    pr.font.size = Pt(8.5)
                    pr.font.bold = True
                    pr.font.color.rgb = RGBColor.from_string("FFFFFF")
                    shade(cell, C.ACCENT)
                ri = 1
            for r_i, row in enumerate(rows):
                for ci, val in enumerate(row):
                    cell = t.rows[ri + r_i].cells[ci]
                    cell.text = ""
                    pr = cell.paragraphs[0].add_run(str(val))
                    pr.font.name = C.FONT_BODY
                    pr.font.size = Pt(8.5)
                    pr.font.color.rgb = RGBColor.from_string(C.INK)
                    if r_i % 2 == 1:
                        shade(cell, "F4F6FA")
            para("", after=6)
        elif k == "sig":
            left, right = b[1], b[2]
            labels = b[3]
            t = d.add_table(rows=5, cols=2)
            t.rows[0].cells[0].text = ""
            for ci, who in enumerate((left, right)):
                pr = t.rows[0].cells[ci].paragraphs[0].add_run(who)
                pr.font.name = C.FONT_BODY
                pr.font.size = Pt(9)
                pr.font.bold = True
                pr.font.color.rgb = RGBColor.from_string(C.INK)
            for r_i, lab in enumerate(labels, start=1):
                for ci in range(2):
                    cell = t.rows[r_i].cells[ci]
                    cell.text = ""
                    pr = cell.paragraphs[0].add_run(lab + " ______________________________")
                    pr.font.name = C.FONT_BODY
                    pr.font.size = Pt(9)
                    pr.font.color.rgb = RGBColor.from_string(C.MUTED)
        elif k == "pagebreak":
            d.add_page_break()
        else:
            raise ValueError("unknown block %r" % (k,))

    d.save(path)
    return path


# --------------------------------------------------------------------------- main
def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--out", default=os.path.join(HERE, "out"))
    a = ap.parse_args(argv)
    os.makedirs(a.out, exist_ok=True)

    print("=" * 78)
    print("  objectale GmbH -> byon gmbh   contract pack   v%s" % C.VERSION)
    print("=" * 78)

    fails = []
    for de_name, en_name, mod in DOCS:
        check(mod.EN, mod.DE, en_name, fails)
    if fails:
        print()
        print("  NOTHING WAS WRITTEN. A bilingual pack whose two versions disagree is worse than")
        print("  a monolingual one, because both get signed.")
        return 1

    made = []
    for de_name, en_name, mod in DOCS:
        for lang, blocks, stem in (("en", mod.EN, en_name + "_EN"),
                                   ("de", mod.DE, de_name + "_DE")):
            p = os.path.join(a.out, "objectale_byon_%s.docx" % stem)
            render(blocks, p, lang)
            made.append((p, len(blocks), len(placeholders(blocks))))

    print()
    for p, nb, nph in made:
        print("  %-62s %3d blocks  %2d field(s) to fill" % (os.path.basename(p), nb, nph))

    every = sorted(set(sum([placeholders(m.EN) + placeholders(m.DE) for _, _, m in DOCS], [])))
    print()
    print("  %d distinct field(s) a human must complete before signature:" % len(every))
    for f in every:
        print("     " + f)
    print()
    print("  %d files in %s" % (len(made), a.out))
    return 0


if __name__ == "__main__":
    sys.exit(main())
